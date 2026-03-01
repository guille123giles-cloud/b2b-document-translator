import streamlit as st
import docx
from pdf2docx import Converter
from deep_translator import GoogleTranslator
import tempfile
import os
import concurrent.futures
import time

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Traductor de Documentos", page_icon="🌍", layout="wide")

def traducir_bloque(texto, origen, destino):
    """Traduce un bloque de texto con reintentos automáticos en caso de micro-cortes."""
    if not texto.strip() or len(texto.strip()) < 2: 
        return texto
    
    for intento in range(3):
        try:
            traductor = GoogleTranslator(source=origen, target=destino)
            return traductor.translate(texto)
        except Exception:
            time.sleep(1.5) # Pausa táctica si hay saturación
    return texto 

def procesar_documento(archivo_subido, origen, destino, barra, estado):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        tmp_pdf.write(archivo_subido.getvalue())
        ruta_pdf = tmp_pdf.name
        
    ruta_docx = ruta_pdf.replace(".pdf", ".docx")
    ruta_docx_out = ruta_pdf.replace(".pdf", "_traducido.docx")

    try:
        estado.info("⚙️ Paso 1/3: Analizando y extrayendo la estructura del documento original...")
        cv = Converter(ruta_pdf)
        cv.convert(ruta_docx, start=0, end=None)
        cv.close()

        doc = docx.Document(ruta_docx)
        
        # Recopilamos todos los textos (párrafos y tablas)
        textos_a_traducir = []
        for parrafo in doc.paragraphs:
            if parrafo.text.strip(): textos_a_traducir.append(parrafo)
            
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        if parrafo.text.strip(): textos_a_traducir.append(parrafo)

        total = len(textos_a_traducir)
        if total == 0:
            estado.warning("⚠️ No se encontró texto legible en el documento.")
            return None

        estado.info(f"🚀 Paso 2/3: Traduciendo {total} fragmentos de texto...")
        
        # 8 hilos: rápido pero estable para la API gratuita
        with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
            futuros = [executor.submit(traducir_bloque, p.text, origen, destino) for p in textos_a_traducir]
            for i, futuro in enumerate(futuros):
                textos_a_traducir[i].text = futuro.result()
                barra.progress((i + 1) / total)

        estado.info("💾 Paso 3/3: Ensamblando el documento final traducido...")
        doc.save(ruta_docx_out)
        
        with open(ruta_docx_out, "rb") as f:
            datos = f.read()
            
        estado.success("✅ ¡Traducción completada con éxito!")
        return datos

    except Exception as e:
        estado.error("❌ Ocurrió un error al procesar la estructura del documento. Asegúrese de que el PDF no esté protegido con contraseña o dañado.")
        return None
    finally:
        for f in [ruta_pdf, ruta_docx, ruta_docx_out]:
            if os.path.exists(f): os.remove(f)

# --- INTERFAZ PARA EL CLIENTE ---
st.title("🌍 Traductor de Documentos")
st.divider()

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("1. Subir Documento")
    archivo = st.file_uploader("Formatos soportados: PDF", type=["pdf"])

with col2:
    st.subheader("2. Opciones de Idioma")
    with st.container(border=True):
        idiomas = {"Auto-detectar": "auto", "Español": "es", "Inglés": "en", "Portugués": "pt", "Francés": "fr"}
        ori_nombre = st.selectbox("Idioma de origen:", list(idiomas.keys()))
        des_nombre = st.selectbox("Traducir al:", [k for k in idiomas.keys() if k != "Auto-detectar"], index=0)
        
        idioma_ori = idiomas[ori_nombre]
        idioma_des = idiomas[des_nombre]

st.divider()

# --- PROCESAMIENTO ---
_, col_btn, _ = st.columns([1, 2, 1])

with col_btn:
    if archivo:
        # Validación de tamaño (ej: aviso si pesa más de 5MB)
        if archivo.size > 5 * 1024 * 1024:
            st.warning("⚠️ El archivo supera los 5MB. El proceso puede demorar más de lo habitual.")
            
        if st.button("🚀 COMENZAR TRADUCCIÓN", use_container_width=True):
            estado = st.empty()
            barra = st.progress(0)
            
            resultado = procesar_documento(archivo, idioma_ori, idioma_des, barra, estado)
            
            if resultado:
                st.download_button(
                    label="⬇️ DESCARGAR DOCUMENTO TRADUCIDO (.DOCX)", 
                    data=resultado, 
                    file_name=f"TRADUCIDO_{archivo.name.replace('.pdf', '')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )